from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "weekly_report_2026-08-26"
ASSET_DIR = OUT_DIR / "assets"
REPORT = OUT_DIR / "TacEx_weekly_report_2026-08-20_to_2026-08-26.docx"

RUNS = [
    {
        "name": "训练1\nEntropy backup + LN",
        "steps": 40740,
        "completed": 1558,
        "success": 837,
        "broken": 197,
        "timeout": 427,
        "other": 97,
        "recent": 54.0,
        "gate": 0.0606,
    },
    {
        "name": "训练2\n关闭backup entropy + LN",
        "steps": 32616,
        "completed": 1258,
        "success": 628,
        "broken": 191,
        "timeout": 368,
        "other": 71,
        "recent": 46.5,
        "gate": 0.0382,
    },
    {
        "name": "训练3\n终止修复 + 平衡回放",
        "steps": 24270,
        "completed": 934,
        "success": 519,
        "broken": 120,
        "timeout": 250,
        "other": 45,
        "recent": 54.7,
        "gate": 0.0267,
    },
]


def rate(run: dict, key: str) -> float:
    return 100.0 * run[key] / run["completed"]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, *, bold: bool = False, color: str | None = None, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(47, 84, 150)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.left_indent = Cm(0.2)
    run = paragraph.add_run("• " + text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.italic = True


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def build_outcome_chart() -> Path:
    labels = ["训练1", "训练2", "训练3"]
    categories = ["成功", "破损", "超时", "其他失败"]
    colors = ["#4C9F70", "#D95F59", "#8795A1", "#E6A23C"]
    values = np.asarray([[rate(run, key) for key in ("success", "broken", "timeout", "other")] for run in RUNS])
    plt.rcParams["font.sans-serif"] = ["Noto Sans CJK JP", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(8.4, 4.8), dpi=180)
    bottom = np.zeros(len(labels))
    for index, (category, color) in enumerate(zip(categories, colors)):
        bars = ax.bar(labels, values[:, index], bottom=bottom, color=color, width=0.62, label=category)
        for bar, value, base in zip(bars, values[:, index], bottom):
            if value >= 6.0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    base + value / 2,
                    f"{value:.1f}%",
                    ha="center",
                    va="center",
                    fontsize=8.5,
                    color="white" if category != "超时" else "#1F2937",
                    fontweight="bold",
                )
        bottom += values[:, index]
    ax.axhline(59.0, color="#17365D", linestyle="--", linewidth=1.2, label="冻结BC成功率59%")
    ax.set_ylim(0, 105)
    ax.set_ylabel("累计回合比例（%）")
    ax.set_title("三条DSRL/SAC训练累计结果（截至2026-08-26）")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.18)
    ax.legend(ncol=3, loc="upper center", bbox_to_anchor=(0.5, -0.10), frameon=False)
    fig.tight_layout()
    path = ASSET_DIR / "dsrl_three_run_outcomes.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def build_trend_chart() -> Path:
    x1 = [37000, 37500, 38000, 38500, 39000, 39500, 40000, 40500]
    y1 = [52.63, 63.16, 55.00, 63.16, 66.67, 52.63, 50.00, 28.57]
    x2 = [29000, 29500, 30000, 30500, 31000, 31500, 32000, 32500]
    y2 = [68.42, 38.10, 63.16, 42.86, 35.29, 30.00, 33.33, 61.11]
    x3 = [20500, 21000, 21500, 22000, 22500, 23000, 23500, 24000]
    y3 = [33.33, 55.00, 52.63, 52.38, 60.00, 52.63, 78.95, 52.63]
    fig, ax = plt.subplots(figsize=(8.4, 4.4), dpi=180)
    ax.plot(x1, y1, marker="o", linewidth=1.6, markersize=4, label="训练1")
    ax.plot(x2, y2, marker="o", linewidth=1.6, markersize=4, label="训练2")
    ax.plot(x3, y3, marker="o", linewidth=1.6, markersize=4, label="训练3")
    ax.axhline(59.0, color="#17365D", linestyle="--", linewidth=1.2, label="冻结BC 59%")
    ax.set_ylim(20, 85)
    ax.set_xlabel("训练步数")
    ax.set_ylabel("最近窗口成功率（%）")
    ax.set_title("最近8个TensorBoard统计窗口的成功率")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(alpha=0.2)
    ax.legend(ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.14), frameon=False)
    fig.tight_layout()
    path = ASSET_DIR / "dsrl_recent_success_trends.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    outcome_chart = build_outcome_chart()
    trend_chart = build_trend_chart()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("TacEx 载玻片抓取项目周报")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(23, 54, 93)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run("周期：2026年8月20日—2026年8月26日")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)

    add_heading(doc, "一、本周工作")
    for text in [
        "围绕冻结Flow Matching BC开展三组门控DSRL/SAC在线训练，分别验证标准entropy backup、关闭backup entropy，以及终止修复与平衡回放方案。",
        "保持BC checkpoint冻结，DSRL输入使用BC双路视觉编码、机器人状态、视觉位姿预测、BC动作摘要和剩余时间，不使用relative_object_pos、object_rot6d或仿真物体真值。",
        "修复SAC将超时后的自动重置观测用于价值自举的问题，并加入终止样本平衡回放、剩余时间观测和晚期无进展惩罚。",
        "截至8月26日，训练1/2/3累计成功率分别为53.7%、49.9%和55.6%；训练3暂时最好，但三条训练均未稳定超过冻结BC的59%基线。",
        "完成代码级复查，确认当前主要瓶颈来自失败类别不可辨识、终止回放只做二分类、单一门控限制全部动作残差，以及宏动作时间尺度下奖励归因不充分。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "二、本周工作具体内容")
    add_heading(doc, "2.1 冻结BC与视觉输入链路", 2)
    for text in [
        "本周DSRL继续使用outputs/lab_pick_flow_bc200_hardyaw_rotation_scratch/best.pt；该BC由100条数据、40轮从零训练获得，正式100回合验证成功59、破损12、掉落2、超时27。",
        "训练和验证范围保持XY连续均匀±10 cm、yaw连续均匀±45°，载玻片破碎阈值保持4 N，本周未修改BC权重或破碎阈值。",
        "确认仿真执行链路使用BC视觉旋转头预测的rot6d，不再用仿真物体朝向覆盖机械臂旋转动作。",
        "BC checkpoint未包含demonstration_mode输入；DSRL观测同样不注入示范模式或成功/失败类别。失败类别仅用于环境终止、奖励和统计。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.2 DSRL/SAC结构与本周改动", 2)
    for text in [
        "残差策略采用4段×10维物理动作残差加1维软门控，共41维；每段修正xyz、rot6d和夹爪宽度，冻结BC仍负责生成完整动作轨迹。",
        "Actor和Critic均使用3层512维MLP，每层加入LayerNorm和ELU。Actor从近零残差初始化，避免训练开始即破坏BC行为。",
        "超时在Gym接口中仍表现为truncated，但写入SAC replay时同时标记为terminal，避免从自动重置后的新回合观测继续进行Q值自举。",
        "训练3加入25%终止样本最低采样比例、0.99宏步折扣、-40超时惩罚、剩余时间观测以及晚期无进展惩罚；BC本体始终冻结。",
        "关闭backup entropy只移除target-Q中的熵项，Actor的熵正则与自动温度学习仍然保留，因此训练2并非完全关闭SAC entropy。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.3 三组训练数据", 2)
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["实验", "步数", "回合", "成功", "破损", "超时", "其他", "近8窗成功"]
    for index, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], "1F4E78")
        set_cell_text(table.rows[0].cells[index], text, bold=True, color="FFFFFF", size=8.3)
    for index, run_data in enumerate(RUNS):
        cells = table.add_row().cells
        values = [
            f"训练{index + 1}",
            f"{run_data['steps']:,}",
            f"{run_data['completed']:,}",
            f"{run_data['success']} / {rate(run_data, 'success'):.1f}%",
            f"{run_data['broken']} / {rate(run_data, 'broken'):.1f}%",
            f"{run_data['timeout']} / {rate(run_data, 'timeout'):.1f}%",
            f"{run_data['other']} / {rate(run_data, 'other'):.1f}%",
            f"{run_data['recent']:.1f}%",
        ]
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, bold=index == 2 and cell is cells[0], size=8.1)
    add_bullet(doc, "训练3累计成功率55.6%，为三组最高；累计超时率26.8%，为三组最低。训练1最近8窗口平均成功率54.0%，训练2为46.5%，训练3为54.7%。")
    add_bullet(doc, "最近窗口通常只覆盖约18—21个回合，方差较大。例如训练3曾出现78.9%的窗口，但下一窗口回落至52.6%，不能据此认定已经收敛。")
    doc.add_picture(str(outcome_chart), width=Cm(14.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图1  三条DSRL/SAC训练的累计回合结果；虚线为冻结BC的59%成功率基线。")

    add_heading(doc, "2.4 训练趋势与诊断", 2)
    doc.add_picture(str(trend_chart), width=Cm(14.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图2  最近8个统计窗口的成功率变化，显示三组训练仍存在明显波动。")
    for text in [
        "三条训练均正常运行，未出现Traceback、CUDA显存溢出或NaN/Inf；Critic loss保持有限，当前不属于数值发散。",
        "训练1累计成功率53.7%，最近8窗口平均54.0%；最新单窗口下降至28.6%，但前序窗口多在50%—66.7%，目前按短期波动处理。",
        "训练2累计成功率49.9%，整体表现最弱；关闭backup entropy没有形成稳定优势，说明仅修改熵备份不足以解决超时和位置失败。",
        "训练3累计成功率55.6%，曾出现较高窗口，但累计破损率仍为12.8%、超时率26.8%，尚未形成持续上升趋势。",
        "门控均值仍较低：训练1约0.0606、训练2约0.0382、训练3约0.0267。策略多数时间只施加很小残差，因此改进幅度受到限制。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "三、代码审查与方案分析")
    for text in [
        "奖励层面：破损、掉落、物体偏离和机械臂越界目前共享普通failure penalty，SAC无法从奖励中明确区分应减小夹取力还是修正抓取位置。",
        "回放层面：TerminalBalancedRandomMemory只根据terminated进行二分类采样，没有分别保证成功、破损、位置失败和超时样本的比例。",
        "动作层面：4段残差共用一个gate；当策略担心某一维引起破损时，降低gate会同时削弱位置、旋转和宽度修正，容易回退到接近BC的策略。",
        "时间尺度层面：一个SAC动作会执行最多16个BC动作，每个动作重复2个物理步。终止奖励只在宏动作末端回传，使精确归因到某段位置或夹爪宽度残差较困难。",
        "评估层面：训练窗口成功率样本量偏小。后续应使用相同的100个固定随机种子评估BC和各DSRL checkpoint，并按验证成功率选择best，而不是按训练reward或最后checkpoint选择。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "四、潜在问题与风险")
    for text in [
        "目前三条DSRL累计成功率均低于冻结BC正式验证的59%，不能宣称残差强化学习已经带来显著改进。",
        "冻结BC基线和在线训练累计数据并非完全相同的一组随机种子；当前横向比较可用于判断趋势，最终结论仍需固定种子的独立100回合评估。",
        "若只提高破损惩罚，策略可能通过减少闭合和残差来降低破损，却把失败转化为超时，造成表面安全但成功率不升。",
        "训练环境仅使用单个视觉仿真实例，单步约5—6秒；完整200k步耗时很长，错误配置会带来较高时间成本。",
        "当前代码工作树包含多项尚未统一提交的BC、环境和DSRL修改；正式上传前需要按功能拆分并完成回归测试。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "五、下周安排")
    for text in [
        "保持4 N破碎阈值和冻结BC不变，新增互斥终止标签：成功、破损、位置失败、超时，并记录掉落、越界、峰值力和各残差维度裁剪比例。",
        "将回放改为按终止类别分层采样，避免终止样本中某一失败类型占比过高；同时保留足够的普通非终止过程样本。",
        "拆分奖励：成功保持+120，位置失败约-25，超时约-35，破损约-50至-60；通过短训练确认不会诱导策略以超时回避破损。",
        "门控采用单次训练内的渐进策略：前期gate上限0.20，积累足够终止样本后逐渐提高到0.30—0.35，同时降低门控惩罚；宽度残差保持较小范围。",
        "在统一环境、奖励和随机范围下重新进行三个seed对照；分别在10k、25k和50k checkpoint使用相同100个随机种子独立验证。",
        "验收目标为成功率至少比冻结BC高10个百分点，同时破损率不高于BC约2个百分点，并要求超时和位置失败持续下降。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "六、数据与附件位置")
    for text in [
        "冻结BC checkpoint：outputs/lab_pick_flow_bc200_hardyaw_rotation_scratch/best.pt",
        "训练1日志：logs/skrl/lab_pick_slide/2026-08-23_13-07-39_sac_torch_dsrl_sac_flow_bc_encoded_physical_gated/",
        "训练2日志：logs/skrl/lab_pick_slide/2026-08-23_18-10-28_sac_torch_dsrl_sac_flow_bc_encoded_physical_gated_no_backup_entropy_ln/",
        "训练3日志：logs/skrl/lab_pick_slide/2026-08-24_17-48-34_sac_torch_dsrl_sac_flow_bc_encoded_physical_gated_terminal_timeouts_balanced_terminal/",
        "本周图表：weekly_report_2026-08-26/assets/dsrl_three_run_outcomes.png、dsrl_recent_success_trends.png",
        "既有任务场景视频：weekly_report_2026-08-19/assets/BC_success_example.mp4（仅用于展示任务场景，不作为本周DSRL定量结果）。",
    ]:
        add_bullet(doc, text)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
